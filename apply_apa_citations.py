# -*- coding: utf-8 -*-
"""Insert APA in-text citations, References before Appendix, update word count, export PDF."""
from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "5004CMD_Project_Report.docx"
PDF = ROOT / "5004CMD_Project_Report.pdf"

CIT_INTRO = (
    " Mobility datasets provide empirical evidence for transport behaviour analysis at national scale "
    "(Bureau of Transportation Statistics, 2024)."
)
CIT_PANDAS = (
    " Pandas provides efficient tabular data manipulation structures widely used in scientific computing "
    "workflows (McKinney, 2010)."
)
CIT_DASK = (
    " Distributed task scheduling enables scalable parallel dataframe processing across cores and clusters "
    "(Dask Development Team, 2024)."
)
CIT_REG = (
    " Linear regression remains a standard baseline method for interpreting relationships between continuous "
    "variables in mobility datasets (James et al., 2021)."
)
CIT_METRICS = (
    " Model evaluation using RMSE and R squared supports interpretation of predictive accuracy and explanatory "
    "strength in statistical learning pipelines (Pedregosa et al., 2011)."
)
CIT_PARALLEL = (
    " Parallel execution performance depends strongly on dataset size partitioning overhead and scheduler "
    "coordination costs (Rocklin, 2015)."
)
CIT_LIM = (
    " Aggregated national mobility datasets may conceal regional variation that affects interpretation of "
    "transport behaviour patterns (Wickham and Grolemund, 2017)."
)

REFERENCES = [
    "Bureau of Transportation Statistics. (2024). Trips by distance dataset. https://data.bts.gov",
    "Dask Development Team. (2024). Dask documentation. https://docs.dask.org",
    "James, G., Witten, D., Hastie, T., Tibshirani, R. (2021). An introduction to statistical learning with applications in Python. Springer.",
    "McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56 to 61.",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., Duchesnay, E. (2011). Scikit learn machine learning in Python. Journal of Machine Learning Research, 12, 2825 to 2830.",
    "Rocklin, M. (2015). Dask parallel computation with blocked algorithms and task scheduling. Proceedings of the 14th Python in Science Conference.",
    "Wickham, H., Grolemund, G. (2017). R for data science. O'Reilly Media.",
]


def append_if_missing(para, suffix: str, marker: str) -> None:
    if marker in para.text:
        return
    para.add_run(suffix)


def insert_references_before_appendix(doc: Document) -> None:
    appendix_p = None
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("Heading") and p.text.strip() == "Appendix":
            appendix_p = p
            break
    if appendix_p is None:
        raise RuntimeError("Appendix heading not found")

    if any("Bureau of Transportation Statistics. (2024). Trips by distance" in p.text for p in doc.paragraphs):
        return

    body = doc.element.body
    blocks: list[tuple[str, str]] = [("h", "References")]
    for ref in REFERENCES:
        blocks.append(("p", ref))

    for kind, text in reversed(blocks):
        if kind == "h":
            new_p = doc.add_heading(text, level=1)
        else:
            new_p = doc.add_paragraph(text)
        el = new_p._element
        body.remove(el)
        appendix_p._element.addprevious(el)


def word_count_excluding_refs_appendix_captions(doc: Document) -> int:
    """Introduction through Conclusion; exclude figure captions; exclude References+Appendix; include timing table."""
    rx_fig = re.compile(r"^Figure\s+\d+\.", re.I)
    collecting = False
    n = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        st = para.style.name if para.style else ""
        if st.startswith("Heading") and t == "Introduction":
            collecting = True
            continue
        if st.startswith("Heading") and t in ("References", "Appendix"):
            break
        if not collecting:
            continue
        if rx_fig.match(t):
            continue
        n += len(re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z]+)?", para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                n += len(re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z]+)?", cell.text))
    return n


def patch_cover_word_count(doc: Document, wc: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for p in doc.paragraphs:
        if "5004CMD Data Mining and Machine Learning" in p.text and "Word count" in p.text:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(
                f"Module\n5004CMD Data Mining and Machine Learning\n\n"
                f"Student ID\n14498572\n\n"
                f"Submission date\n30 March 2026\n\n"
                f"Word count\n{wc}"
            )
            return
    raise RuntimeError("Cover metadata paragraph not found.")


def export_pdf_word() -> None:
    dx = DOCX.resolve().as_posix().replace("'", "''")
    pdf = PDF.resolve().as_posix().replace("'", "''")
    ps = rf"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $d = $word.Documents.Open('{dx}')
  $d.SaveAs2([ref]'{pdf}', [ref]17) | Out-Null
  $d.Close()
}} finally {{
  $word.Quit()
  [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
}}
"""
    r = subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps],
        capture_output=True,
        text=True,
    )
    if r.returncode != 0 or not PDF.is_file():
        print("PDF export stderr:", r.stderr, file=sys.stderr)


def main() -> int:
    if not DOCX.is_file():
        print("Run build_5004cmd_report.py first.", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))

    for p in doc.paragraphs:
        t = p.text
        if t.startswith("Mobility datasets quantify"):
            append_if_missing(p, CIT_INTRO, "Bureau of Transportation Statistics, 2024")
        elif t.startswith("Pandas supplies the baseline"):
            append_if_missing(p, CIT_PANDAS, "McKinney, 2010")
        elif t.startswith("Dask partitions frames"):
            append_if_missing(p, CIT_DASK, "Dask Development Team, 2024")
        elif "violations can mislead" in t and "Figure 5" in t:
            append_if_missing(p, CIT_REG, "James et al., 2021")
        elif t.startswith("R2 equals 0.946 implies"):
            append_if_missing(p, CIT_METRICS, "Pedregosa et al., 2011")
        elif "moderate data scale" in t and "recorded coursework run" in t:
            append_if_missing(p, CIT_PARALLEL, "Rocklin, 2015")
        elif "residual variation unexplained" in t and "National aggregation hides" in t:
            append_if_missing(p, CIT_LIM, "Wickham and Grolemund, 2017")

    insert_references_before_appendix(doc)

    wc = word_count_excluding_refs_appendix_captions(doc)
    patch_cover_word_count(doc, wc)

    doc.save(str(DOCX))
    print("Saved", DOCX, "| Word count (excl. refs, appendix, captions):", wc)

    export_pdf_word()
    if PDF.is_file():
        print("Saved", PDF)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
