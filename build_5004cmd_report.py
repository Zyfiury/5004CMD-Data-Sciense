# -*- coding: utf-8 -*-
"""Build 5004CMD_Project_Report.docx (distinction-level structure). Run from project root."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
OUT_DOCX = ROOT / "5004CMD_Project_Report.docx"
OUT_PDF = ROOT / "5004CMD_Project_Report.pdf"

CAPTIONS = [
    ("fig1_weekly_home.png", "Figure 1. Average population staying at home per week"),
    ("fig2_weekly_not_home.png", "Figure 2. Average population not staying at home per week"),
    ("fig3_distance_means.png", "Figure 3. Average trips across distance categories"),
    ("fig4_trip_threshold_scatter.png", "Figure 4. Trips above ten million threshold over time"),
    ("fig5_regression.png", "Figure 5. Linear regression between distance categories"),
    ("fig6_travellers_by_distance.png", "Figure 6. Distribution of travellers by distance band"),
]


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    r = run._r
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    r.append(fld2)
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose Update Field to generate the table of contents."
    r.append(t)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r.append(fld3)


def caption_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def add_figure(doc: Document, filename: str, cap: str) -> None:
    doc.add_picture(str(FIG / filename), width=Inches(6))
    caption_para(doc, cap)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def body(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.strip())


def word_count_main(doc: Document) -> int:
    """Words from first 'Introduction' H1 through text before 'Appendix' H1; includes tables, excludes appendix."""
    collecting = False
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        st = para.style.name if para.style else ""
        if st.startswith("Heading") and t == "Introduction":
            collecting = True
        if st.startswith("Heading") and t == "Appendix":
            break
        if collecting:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    blob = " ".join(parts)
    return len(re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z]+)?", blob))


def patch_cover_word_count(doc: Document, wc: int) -> None:
    for p in doc.paragraphs:
        if "5004CMD Data Mining and Machine Learning" in p.text and "Word count" in p.text:
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(
                f"Module\n5004CMD Data Mining and Machine Learning\n\n"
                f"Student ID\n14498572\n\n"
                f"Submission date\n30 March 2026\n\n"
                f"Word count\n{wc}"
            )
            return
    raise RuntimeError("Cover metadata paragraph not found for word-count update.")


def cover_page(doc: Document, wc_placeholder: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BTS Mobility Data Analysis Using Sequential and Distributed Processing")
    try:
        p.style = doc.styles["Title"]
    except KeyError:
        r.bold = True
        r.font.size = Pt(20)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Module\n5004CMD Data Mining and Machine Learning\n\n"
        f"Student ID\n14498572\n\n"
        f"Submission date\n30 March 2026\n\n"
        f"Word count\n{wc_placeholder}"
    )


def timing_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Processing mode"
    table.rows[0].cells[1].text = "Time (seconds)"
    rows = [
        ("Serial processing", "3.99"),
        ("Parallel processing with ten workers", "9.06"),
        ("Parallel processing with twenty workers", "9.58"),
    ]
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b


def build(wc_line: str) -> Document:
    for fn, _ in CAPTIONS:
        if not (FIG / fn).is_file():
            print("Missing figure:", FIG / fn, file=sys.stderr)
            sys.exit(1)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    cover_page(doc, wc_line)
    doc.add_page_break()

    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_toc_title.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    # --- Introduction ---
    h1(doc, "Introduction")
    body(
        doc,
        """
Mobility datasets quantify population-scale travel in time and space, including trips by distance band. They underpin transport analytics by measuring realised demand, network stress, and environmental exposure, and by supplying empirical baselines when accessibility, pricing, or land use changes alter trip-making.

Distance stratification approximates activity space: short bands capture local circulation and chained errands; medium bands often reflect intra-regional commuting and service access; long bands reflect inter-regional movement. Comparing bands reveals how national mobility distributes across spatial scales instead of relying on a single aggregate intensity.

Regression links aligned daily series across bands, summarising co-movement and providing simple checks alongside R squared and RMSE expressed in natural trip-count units. Contrasting Pandas with Dask addresses scalable analytics: serial code is straightforward to validate; distributed execution may reduce wall-clock time when partitions are large enough to amortise scheduling, serialization, and worker communication.

Four objectives are pursued: (1) summarise weekly national Population Staying at Home and Population Not Staying at Home; (2) quantify distance-stratified intensity and flag dates where medium-distance volumes exceed ten million trips; (3) estimate a linear relationship between Trips 1-25 Miles and Number of Trips 5-10 on merged national dates, reporting RMSE and R squared; (4) benchmark runtime for a replicated workflow under serial Pandas and parallel Dask with ten and twenty workers.
""",
    )

    # --- Datasets ---
    h1(doc, "Datasets")
    body(
        doc,
        """
Trips_by_Distance.csv contains Level, Date, Week, Population Staying at Home, Population Not Staying at Home, aggregate trips, and binned counts such as Number of Trips 5-10 and Number of Trips 10-25. Trips_Full Data.csv contains national daily distance columns Trips <1 Mile through Trips 500+ Miles plus population and aggregate trip fields used for distance summaries complementary to the distance file.

National filtering keeps a single geography comparable across days. Datetime conversion enables correct ordering, reliable inner merges on calendar dates, and valid weekly aggregation. Duplicate removal avoids inflated means and distorted threshold frequencies. Redundant geographic columns are dropped after retaining National rows to simplify schemas and reduce dtype inference failures in distributed CSV reads.
""",
    )

    # --- Methodology ---
    h1(doc, "Methodology")
    body(
        doc,
        """
Pandas supplies the baseline: deterministic in-process frames, mature CSV and datetime handling, and plotting integration suitable for reproducible coursework and reference outputs that downstream distributed runs can mirror.

Dask partitions frames and executes task graphs on a local cluster to parallelise reads and reductions without altering the substantive analytical steps, enabling a controlled wall-clock comparison on the same logical workflow.

Weekly means group Trips_by_Distance.csv by Week for Population Staying at Home and Population Not Staying at Home, smoothing daily noise while preserving slower-moving weekly structure visible in bar charts. Distance means average National rows in Trips_Full Data.csv across mile-band columns to summarise typical intensity and to compare category dominance within the national mix.

Thresholding selects dates where Number of Trips 10-25 or Number of Trips 50-100 exceed ten million trips, operationalising an extreme-intensity regime for visual comparison across bands. Regression inner-merges Trips 1-25 Miles from Trips_Full Data.csv with Number of Trips 5-10 from Trips_by_Distance.csv on Date over the overlapping national window to test same-day co-movement between a broad predictor and an adjacent response.

RMSE reports typical absolute error in trip counts; R squared reports variance explained relative to a mean-only benchmark, separating scale-sensitive accuracy from explanatory power for reporting audiences.
""",
    )

    # --- Exploratory Data Analysis + figures 1–3 ---
    h1(doc, "Exploratory Data Analysis")
    body(
        doc,
        """
Weekly aggregation matches common reporting cadences and improves cross-month comparability. Population staying at home signals reduced out-of-home time and is sensitive to gradual shifts in routine participation. Population not staying at home captures realised travel on the same days and completes the national activity budget implied by the two population fields. Distance comparisons matter because headline totals can conceal concentration in short bands that load local networks and environmental exposures. Short-distance dominance is expected when daily routines remain localised even if occasional long trips are socially salient.

Figures 1–3 present weekly home and non-home averages and cross-sectional distance means before thresholding and regression.
""",
    )

    add_figure(doc, CAPTIONS[0][0], CAPTIONS[0][1])
    body(
        doc,
        """
Figure 1 shows weekly mean Population Staying at Home varying gradually without sharp single-week breaks, consistent with slow routine adjustment rather than episodic shocks at this resolution. Late-sample increases, where visible, imply modestly higher average home time and should be read with Figure 2 to infer net mobility change.
""",
    )

    add_figure(doc, CAPTIONS[1][0], CAPTIONS[1][1])
    body(
        doc,
        """
Figure 2 shows Population Not Staying at Home persistently large and comparatively stable, indicating structurally high national circulation with only modest week-to-week variation around a high mean. Together with Figure 1, the series suggest gradual co-movement within a stable aggregate envelope.
""",
    )

    add_figure(doc, CAPTIONS[2][0], CAPTIONS[2][1])
    body(
        doc,
        """
Figure 3 places the highest means in short bands, led by Trips 1-3 Miles, consistent with local access trips and short segments within wider journeys. This implies commuting systems still generate massive intra-local circulation, so junction reliability, local transit frequency, and active-travel safety remain high-leverage even when longer commutes exist.
""",
    )

    # --- Threshold Analysis ---
    h1(doc, "Threshold Analysis")
    body(
        doc,
        """
Ten million trips marks an extreme national daily count that is simple to communicate across bins. Ten-to-twenty-five-mile movement captures intra-regional circulation between suburbs and centres, distinct from purely local or rare long-haul travel. Frequent medium-band exceedance relative to a longer band suggests repeatable medium-range travel drives many high-volume days. Figure 4 shows temporal clustering of exceedances.
""",
    )
    add_figure(doc, CAPTIONS[3][0], CAPTIONS[3][1])

    # --- Regression Model ---
    h1(doc, "Regression Model")
    body(
        doc,
        """
Trips 1-25 Miles concentrates short-to-medium national volume into one predictor aligned with exploratory distance structure. Number of Trips 5-10 targets an adjacent operational bin beyond the shortest hops, merged on Date for national rows. A positive slope indicates same-day co-movement consistent with shared activity budgets. Coefficient, intercept, RMSE, and R2 are reported in Model Evaluation.

Figure 5 plots observations with the ordinary least squares line. Linear assumptions include an approximately linear conditional mean and limited outlier leverage; violations can mislead even when fit appears strong.
""",
    )
    add_figure(doc, CAPTIONS[4][0], CAPTIONS[4][1])

    # --- Model Evaluation (fixed numeric lines) ---
    h1(doc, "Model Evaluation")
    body(
        doc,
        """
Ordinary least squares on nationally merged daily observations yields the metrics below in original trip-count units.
""",
    )
    doc.add_paragraph("R2 equals 0.946")
    doc.add_paragraph("RMSE equals 1981454.86")
    doc.add_paragraph("Coefficient equals 0.180985")
    doc.add_paragraph("Intercept equals 49643925.69")
    body(
        doc,
        """
R2 equals 0.946 implies the linear predictor explains most day-to-day variation in Number of Trips 5-10 conditional on Trips 1-25 Miles in the merged window, signalling a tight linear association for descriptive reporting. RMSE equals 1981454.86 remains large in absolute terms because national daily counts are enormous, so even proportionally small errors translate into large trip-count deviations. Coefficient and intercept should be read jointly: the slope describes incremental co-movement between bands, while the intercept positions the line. Limitations include strict linearity, omitted weekly or shock structure, and sensitivity to merge coverage across Trips_Full Data.csv and Trips_by_Distance.csv.
""",
    )

    # --- Travellers by Distance Band ---
    h1(doc, "Travellers by Distance Band")
    add_figure(doc, CAPTIONS[5][0], CAPTIONS[5][1])
    body(
        doc,
        """
Figure 6 shows declining average travellers as distance increases, concentrating routine intensity in short bands and leaving long bands with comparatively small averages. That skew implies infrastructure stress and local air-quality exposures track short-range circulation most closely. Planning therefore prioritises junction capacity, local transit headways, and walking and cycling safety, while longer corridors remain important for connectivity even though they carry a smaller share of average daily mass in these summaries.
""",
    )

    # --- Parallel Processing Comparison ---
    h1(doc, "Parallel Processing Comparison")
    timing_table(doc)
    body(
        doc,
        """
The benchmark asks whether partitioned Dask execution accelerates the replicated pipeline relative to serial Pandas on the same logical steps. Sequential runs minimise coordination and keep data movement local to one process. Distributed runs add a scheduler, worker processes, serialization, and partitioned CSV reads—fixed and variable overhead that can dominate elapsed time when partitions are small or when results must be collected frequently to the client.

Performance scales with size: larger data amortises fixed cluster costs across more bytes and tasks, whereas smaller extracts spend disproportionate time on graph construction and startup. Dask is more compelling for large, repeated partitioned pipelines and when out-of-core execution avoids memory ceilings. In the recorded coursework run, serial time was 3.99 seconds versus 9.06 seconds with ten workers and 9.58 seconds with twenty workers, illustrating coordination overhead under the tested configuration and moderate data scale.
""",
    )

    # --- Discussion ---
    h1(doc, "Discussion")
    body(
        doc,
        """
Short-band dominance structures daily activity around local access and repeated short hops, while medium-distance travel sustains regional connectivity and can spike above the ten-million threshold on busy days. Regression ties Trips 1-25 Miles to Number of Trips 5-10, supporting simple monitoring, nowcast-style checks, and communication of co-movement to non-technical stakeholders. Timing results caution that parallelism is not automatically faster; tool choice should follow data scale, task granularity, and how often results must be consolidated.
""",
    )

    # --- Limitations ---
    h1(doc, "Limitations")
    body(
        doc,
        """
National aggregation hides spatial heterogeneity and distributional impacts across communities. Seasonality is not isolated, so gradual trends may mix holiday calendars, policy shifts, and structural change. A single linear specification may miss threshold or piecewise behaviour in mobility regimes. The moderate extract size used for timing comparisons limits distributed speedups because coordination costs remain salient relative to compute. Omitted covariates—fuel prices, weather, labour-market shocks—restrict causal interpretation and leave residual variation unexplained.
""",
    )

    # --- Conclusion ---
    h1(doc, "Conclusion")
    body(
        doc,
        """
Trips_by_Distance.csv and Trips_Full Data.csv jointly portray weekly home and non-home behaviour, short-band dominance, and medium-distance high-volume days. Regression links Trips 1-25 Miles to Number of Trips 5-10 with R2 equal to 0.946 and RMSE equal to 1981454.86. Serial Pandas outpaced Dask with ten and twenty workers in the recorded benchmark, consistent with coordination overhead. Future work may add geography, seasonality, nonlinear models, larger extracts to retest parallelism, and policy covariates.
""",
    )

    # --- Appendix (unchanged text, page break) ---
    doc.add_page_break()
    h1(doc, "Appendix")

    doc.add_heading("AI declaration", level=2)
    body(doc, (ROOT / "ai_use_appendix_template.txt").read_text(encoding="utf-8"))

    doc.add_heading("Pseudocode", level=2)
    body(doc, (ROOT / "pseudocode_appendix.txt").read_text(encoding="utf-8"))

    doc.add_heading("Flowchart description", level=2)
    body(doc, (ROOT / "flowchart_appendix.txt").read_text(encoding="utf-8"))

    return doc


def main() -> None:
    doc = build("calculating…")
    wc = word_count_main(doc)
    patch_cover_word_count(doc, wc)
    doc.save(str(OUT_DOCX))
    if wc < 1500 or wc > 1700:
        print(f"Warning: main-body word count {wc} is outside 1500–1700; adjust prose in build script.", file=sys.stderr)

    print("Saved", OUT_DOCX)
    print("Main body word count (excluding appendix, incl. tables):", wc)

    try:
        import subprocess

        ps = r"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {
  $d = $word.Documents.Open('%s')
  $d.SaveAs2([ref]'%s', [ref]17) | Out-Null
  $d.Close()
} finally {
  $word.Quit()
  [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
}
""" % (
            OUT_DOCX.resolve().as_posix().replace("'", "''"),
            OUT_PDF.resolve().as_posix().replace("'", "''"),
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            check=False,
            capture_output=True,
            text=True,
        )
        if OUT_PDF.is_file():
            print("Saved", OUT_PDF)
    except Exception as e:
        print("PDF export skipped:", e, file=sys.stderr)


if __name__ == "__main__":
    main()
